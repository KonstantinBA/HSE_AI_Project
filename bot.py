import asyncio
import logging
import os
from docx import Document

from dotenv import load_dotenv
from datetime import datetime
from typing import Any, Awaitable, Callable, Dict

import nest_asyncio
import aiosqlite
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from pytz import timezone

from aiogram import Bot, Dispatcher, BaseMiddleware, types
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ParseMode
from aiogram.filters import CommandStart, Command
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.fsm.storage.memory import MemoryStorage
from aiogram.types import (
    BotCommand, 
    CallbackQuery, 
    InlineKeyboardButton,
    InlineKeyboardMarkup,
    KeyboardButton,
    Message,
    ReplyKeyboardMarkup,
    ReplyKeyboardRemove,
    FSInputFile
)

# ================= LANGCHAIN –ò LANGGRAPH =================
from langchain_core.messages import HumanMessage
from langchain_gigachat.chat_models import GigaChat
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langgraph.checkpoint.memory import MemorySaver
from langgraph.graph import START, MessagesState, StateGraph
from datetime import datetime, timedelta, timezone
# –ó–∞–≥—Ä—É–∂–∞–µ–º –ø–µ—Ä–µ–º–µ–Ω–Ω—ã–µ –∏–∑ —Ñ–∞–π–ª–∞ .env
load_dotenv()

API_TOKEN = os.getenv('API_TOKEN')
API_KEY = os.getenv('GIGACHAT_KEY')
DB_PATH = "database/users.db"
gmt_plus_3 = timezone(timedelta(hours=3))
print(gmt_plus_3)
# --- Logging configuration ---
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è GigaChat (–µ—Å–ª–∏ –Ω–µ–æ–±—Ö–æ–¥–∏–º–æ, –∑–∞–º–µ–Ω–∏—Ç–µ –ø–∞—Ä–∞–º–µ—Ç—Ä—ã)
model = GigaChat(
    credentials=API_KEY,
    scope="GIGACHAT_API_PERS",
    model="GigaChat",
    verify_ssl_certs=False,
)

# –ü—Ä–æ–º–ø—Ç –¥–ª—è –≤–∑–∞–∏–º–æ–¥–µ–π—Å—Ç–≤–∏—è —Å –º–æ–¥–µ–ª—å—é
prompt = ChatPromptTemplate.from_messages(
    [
        ("system", "–¢—ã ‚Äî –ø–æ–º–æ—â–Ω–∏–∫. –û—Ç–≤–µ—á–∞–π –∫—Ä–∞—Ç–∫–æ, –Ω–æ –ø–æ–ª–µ–∑–Ω–æ."),
        MessagesPlaceholder(variable_name="messages"),
    ]
)

workflow = StateGraph(state_schema=MessagesState)

# –ê—Å–∏–Ω—Ö—Ä–æ–Ω–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è –¥–ª—è –≤—ã–∑–æ–≤–∞ –º–æ–¥–µ–ª–∏
async def call_model(state: MessagesState):
    # –°–æ–∑–¥–∞–µ–º —Ü–µ–ø–æ—á–∫—É: —Å–Ω–∞—á–∞–ª–∞ –∏—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è prompt, –∑–∞—Ç–µ–º –º–æ–¥–µ–ª—å
    chain = prompt | model
    response = await chain.ainvoke(state)
    return {"messages": response}

# –î–æ–±–∞–≤–ª—è–µ–º –≤–µ—Ä—à–∏–Ω—É –≥—Ä–∞—Ñ–∞
workflow.add_edge(START, "model")
workflow.add_node("model", call_model)

# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –ø–µ—Ä—Å–∏—Å—Ç–µ–Ω—Ç–Ω–æ—Å—Ç—å —á–µ—Ä–µ–∑ MemorySaver
memory = MemorySaver()

# –ö–æ–º–ø–∏–ª–∏—Ä—É–µ–º –≥—Ä–∞—Ñ, –ø–æ–ª—É—á–∞—è –ø—Ä–∏–ª–æ–∂–µ–Ω–∏–µ –¥–ª—è –≤—ã–∑–æ–≤–∞ –º–æ–¥–µ–ª–∏
app = workflow.compile(checkpointer=memory)

# --- Scheduler initialization ---
scheduler = AsyncIOScheduler()

# --- Bot and Dispatcher initialization ---
bot = Bot(
    token=API_TOKEN,
    default=DefaultBotProperties(parse_mode=ParseMode.HTML)
)
storage = MemoryStorage()
dp = Dispatcher(storage=storage)

# --- States ---
class DialogForm(StatesGroup):
    in_dialog = State()

class RegistrationForm(StatesGroup):
    name = State()
    age = State()
    email = State()

class DiaryForm(StatesGroup):
    situation = State()
    thought = State()
    emotion = State()
    reaction = State()

class ReminderForm(StatesGroup):
    time = State()

# --- Keyboards ---
main_menu = ReplyKeyboardMarkup(
    keyboard=[
        [KeyboardButton(text="–î–æ–±–∞–≤–∏—Ç—å –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫"), KeyboardButton(text="–ü–æ–ª—É—á–∏—Ç—å —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—é")],
        [KeyboardButton(text="–ü–æ—Å–º–æ—Ç—Ä–µ—Ç—å –¥–Ω–µ–≤–Ω–∏–∫"), KeyboardButton(text="–≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –¥–Ω–µ–≤–Ω–∏–∫")],
        [KeyboardButton(text="–û—Å—Ç–∞–≤–∏—Ç—å –æ—Ç–∑—ã–≤"), KeyboardButton(text="–ù–∞—Å—Ç—Ä–æ–π–∫–∏")]
    ],
    resize_keyboard=True,
    one_time_keyboard=True
)

settings_menu = InlineKeyboardMarkup(
    inline_keyboard=[
        [
            InlineKeyboardButton(
                text="\u2705 –ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è: –í–∫–ª—é—á–∏—Ç—å", 
                callback_data="toggle_reminder_on"
            ),
            InlineKeyboardButton(
                text="\u274C –ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è: –í—ã–∫–ª—é—á–∏—Ç—å", 
                callback_data="toggle_reminder_off"
            )
        ],
        [
            InlineKeyboardButton(
                text="–£—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –≤—Ä–µ–º—è –Ω–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è", 
                callback_data="set_reminder_time"
            )
        ]
    ]
)

dialog_buttons = InlineKeyboardMarkup(inline_keyboard=[
    [InlineKeyboardButton(text="–ü—Ä–æ–¥–æ–ª–∂–∏—Ç—å –¥–∏–∞–ª–æ–≥", callback_data="continue_dialog")],
    [InlineKeyboardButton(text="–ó–∞–≤–µ—Ä—à–∏—Ç—å –¥–∏–∞–ª–æ–≥", callback_data="end_dialog")]
])

# --- Middleware for registration check ---
class RegistrationMiddleware(BaseMiddleware):
    async def __call__(
        self,
        handler: Callable[[types.TelegramObject, Dict[str, Any]], Awaitable[Any]],
        event: types.TelegramObject,
        data: Dict[str, Any]
    ) -> Any:
        print(isinstance(event, Message))
        if isinstance(event, Message) and event.text != '/start':
            user_id = event.chat.id
            fsm_context: FSMContext = data["state"]
            state = await fsm_context.get_state()

            # –ï—Å–ª–∏ –∏–¥—ë—Ç –ø—Ä–æ—Ü–µ—Å—Å —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏–∏, –ø—Ä–æ–ø—É—Å–∫–∞–µ–º
            if state and state.startswith("RegistrationForm:"):
                return await handler(event, data)

            # –ò–Ω–∞—á–µ –ø—Ä–æ–≤–µ—Ä—è–µ–º, –µ—Å—Ç—å –ª–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å –≤ –ë–î
            async with aiosqlite.connect(DB_PATH) as db:
                async with db.execute(
                    "SELECT id FROM users WHERE id = ?", (user_id,)
                ) as cursor:
                    if await cursor.fetchone() is None:
                        await bot.send_message(
                            chat_id=user_id,
                            text="–í—ã –Ω–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã! –ó–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä—É–π—Ç–µ—Å—å, –∏—Å–ø–æ–ª—å–∑—É—è –∫–æ–º–∞–Ω–¥—É /start."
                        )
                        return

        return await handler(event, data)

# --- Reminder Function ---
async def send_reminders():
    current_time = datetime.now(gmt_plus_3).strftime("%H:%M")
    current_date = datetime.now(gmt_plus_3).date()

    print(f"[DEBUG] Current time: {current_time}, Current date: {current_date}")

    async with aiosqlite.connect(DB_PATH) as db:
        async with db.execute(
            """
            SELECT user_id, last_sent_date
            FROM reminders
            WHERE enabled = 1 AND time = ?
            """,
            (current_time,)
        ) as cursor:
            users = await cursor.fetchall()
            for user_id, last_sent_date in users:
                # –ï—Å–ª–∏ —É–∂–µ –æ—Ç–ø—Ä–∞–≤–ª—è–ª–∏ —Å–µ–≥–æ–¥–Ω—è
                if last_sent_date == str(current_date):
                    continue

                try:
                    await bot.send_message(
                        user_id,
                        "–ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏–µ: –ù–µ –∑–∞–±—É–¥—å—Ç–µ –¥–æ–±–∞–≤–∏—Ç—å –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫!"
                    )
                    # –û–±–Ω–æ–≤–ª—è–µ–º –¥–∞—Ç—É –ø–æ—Å–ª–µ–¥–Ω–µ–π –æ—Ç–ø—Ä–∞–≤–∫–∏
                    await db.execute(
                        """
                        UPDATE reminders
                        SET last_sent_date = ?
                        WHERE user_id = ?
                        """,
                        (current_date, user_id)
                    )
                    await db.commit()
                except Exception as e:
                    logger.error(
                        f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ—Ç–ø—Ä–∞–≤–∏—Ç—å –Ω–∞–ø–æ–º–∏–Ω–∞–Ω–∏–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é {user_id}: {e}"
                    )

# --- Utility Functions ---
async def get_last_diary_entry(user_id: int):
    async with aiosqlite.connect(DB_PATH) as db:
        async with db.execute(
            """
            SELECT id, situation, thought, emotion, reaction
            FROM diary
            WHERE user_id = ?
            ORDER BY created_at DESC
            LIMIT 1
            """,
            (user_id,)
        ) as cursor:
            entry = await cursor.fetchone()
            if entry:
                return entry[1], entry[2], entry[3], entry[4], entry[0]
            return None

# --- Register Middleware ---
dp.update.middleware.register(RegistrationMiddleware())

# --- Handlers ---
@dp.message(CommandStart())
async def cmd_start(message: Message, state: FSMContext):
    user_id = message.from_user.id
    async with aiosqlite.connect(DB_PATH) as db:
        async with db.execute(
            "SELECT id, name FROM users WHERE id = ?", (user_id,)
        ) as cursor:
            user = await cursor.fetchone()
            # –ï—Å–ª–∏ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –Ω–µ—Ç, –Ω–∞—á–∏–Ω–∞–µ–º —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—é
            if user is None:
                await state.set_state(RegistrationForm.name)
                await message.answer("–ù–µ–æ–±—Ö–æ–¥–∏–º–æ –ø—Ä–æ–π—Ç–∏ —Ä–µ–≥–∏—Å—Ç—Ä–∞—Ü–∏—é!")
                await message.answer("–í–≤–µ–¥–∏—Ç–µ –≤–∞—à–µ –∏–º—è:")
            else:
                await message.answer(
                    f"–ü—Ä–∏–≤–µ—Ç, {user[1]}! –í—ã —É–∂–µ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã!",
                    reply_markup=main_menu
                )

@dp.message(RegistrationForm.name)
async def process_name(message: Message, state: FSMContext):
    await state.update_data(name=message.text)
    await state.set_state(RegistrationForm.age)
    await message.answer("–í–≤–µ–¥–∏—Ç–µ –≤–∞—à –≤–æ–∑—Ä–∞—Å—Ç:")

@dp.message(RegistrationForm.age)
async def process_age(message: Message, state: FSMContext):
    if not message.text.isdigit():
        await message.answer("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –≤–≤–µ–¥–∏—Ç–µ –∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–π –≤–æ–∑—Ä–∞—Å—Ç:")
        return
    await state.update_data(age=int(message.text))
    await state.set_state(RegistrationForm.email)
    await message.answer("–í–≤–µ–¥–∏—Ç–µ –≤–∞—à email:")

@dp.message(RegistrationForm.email)
async def process_email(message: Message, state: FSMContext):
    await state.update_data(email=message.text)
    user_data = await state.get_data()

    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            """
            INSERT INTO users (id, name, age, email)
            VALUES (?, ?, ?, ?)
            """,
            (message.from_user.id, user_data["name"], user_data["age"], user_data["email"])
        )
        await db.commit()

    await state.clear()
    await message.answer(
        f"{user_data['name']}, –≤—ã —É—Å–ø–µ—à–Ω–æ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞–Ω—ã!",
        reply_markup=main_menu
    )

@dp.message(Command(commands=["new_entry"]))
async def cmd_new_entry(message: Message, state: FSMContext):
    await state.set_state(DiaryForm.situation)
    await message.answer("–û–ø–∏—à–∏—Ç–µ —Å–∏—Ç—É–∞—Ü–∏—é, –∫–æ—Ç–æ—Ä–∞—è –ø—Ä–æ–∏–∑–æ—à–ª–∞:")

@dp.message(lambda m: m.text == "–î–æ–±–∞–≤–∏—Ç—å –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫")
async def handle_menu_new_entry(message: Message, state: FSMContext):
    await state.set_state(DiaryForm.situation)
    await message.answer("–û–ø–∏—à–∏—Ç–µ —Å–∏—Ç—É–∞—Ü–∏—é, –∫–æ—Ç–æ—Ä–∞—è –ø—Ä–æ–∏–∑–æ—à–ª–∞:", reply_markup=ReplyKeyboardRemove())

@dp.message(DiaryForm.situation)
async def process_situation(message: Message, state: FSMContext):
    await state.update_data(situation=message.text)
    await state.set_state(DiaryForm.thought)
    await message.answer("–ö–∞–∫–∞—è –º—ã—Å–ª—å —É –≤–∞—Å –≤–æ–∑–Ω–∏–∫–ª–∞?", reply_markup=ReplyKeyboardRemove())

@dp.message(DiaryForm.thought)
async def process_thought(message: Message, state: FSMContext):
    await state.update_data(thought=message.text)
    await state.set_state(DiaryForm.emotion)
    await message.answer("–ö–∞–∫–∏–µ —ç–º–æ—Ü–∏–∏ –≤—ã –∏—Å–ø—ã—Ç–∞–ª–∏?", reply_markup=ReplyKeyboardRemove())

@dp.message(DiaryForm.emotion)
async def process_emotion(message: Message, state: FSMContext):
    await state.update_data(emotion=message.text)
    await state.set_state(DiaryForm.reaction)
    await message.answer("–ö–∞–∫–∞—è –±—ã–ª–∞ –≤–∞—à–∞ —Ä–µ–∞–∫—Ü–∏—è?", reply_markup=ReplyKeyboardRemove())

@dp.message(DiaryForm.reaction)
async def process_reaction(message: Message, state: FSMContext):
    await state.update_data(reaction=message.text)
    user_data = await state.get_data()

    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            """
            INSERT INTO diary (user_id, situation, thought, emotion, reaction)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                message.from_user.id,
                user_data["situation"],
                user_data["thought"],
                user_data["emotion"],
                user_data["reaction"]
            )
        )
        await db.commit()

    await state.clear()
    await message.answer("–ó–∞–ø–∏—Å—å —É—Å–ø–µ—à–Ω–æ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∞!", reply_markup=main_menu)

# --- –û–±—Ä–∞–±–æ—Ç—á–∏–∫ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–π ---
@dp.message(Command(commands=["get_recommendation"]))
@dp.message(lambda m: m.text == "–ü–æ–ª—É—á–∏—Ç—å —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—é")
async def handle_menu_get_recommendation(message: Message):
    user_id = str(message.from_user.id)

    await message.answer("–°–µ–π—á–∞—Å –ø–æ—Å–º–æ—Ç—Ä–∏–º, –ø–æ–¥–æ–∂–¥–∏—Ç–µ...")

    # –ü–æ–ª—É—á–µ–Ω–∏–µ –ø–æ—Å–ª–µ–¥–Ω–µ–π –∑–∞–ø–∏—Å–∏ (—Ñ—É–Ω–∫—Ü–∏—è –ø–æ–ª—É—á–µ–Ω–∏—è –∑–∞–ø–∏—Å–∏ –∏–∑ –ë–î)
    last_entry = await get_last_diary_entry(int(user_id))
    if not last_entry:
        await message.answer("–ù–µ—Ç –∑–∞–ø–∏—Å–µ–π –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞. –î–æ–±–∞–≤—å—Ç–µ –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫.")
        return

    situation, thought, emotion, reaction, entry_id = last_entry

    # –§–æ—Ä–º–∏—Ä—É–µ–º —Ü–µ–ø–æ—á–∫—É —Å–æ–æ–±—â–µ–Ω–∏–π
    input_messages = [
        HumanMessage(
            content=f"""–¢—ã ‚Äî –∫–≤–∞–ª–∏—Ñ–∏—Ü–∏—Ä–æ–≤–∞–Ω–Ω—ã–π –ø—Å–∏—Ö–æ–ª–æ–≥ —Å –≥–ª—É–±–æ–∫–∏–º –ø–æ–Ω–∏–º–∞–Ω–∏–µ–º –∫–æ–≥–Ω–∏—Ç–∏–≤–Ω–æ-–ø–æ–≤–µ–¥–µ–Ω—á–µ—Å–∫–æ–π —Ç–µ—Ä–∞–ø–∏–∏ –∏ —ç–º–æ—Ü–∏–æ–Ω–∞–ª—å–Ω–æ–≥–æ –∏–Ω—Ç–µ–ª–ª–µ–∫—Ç–∞.
                –¢–≤–æ—è –∑–∞–¥–∞—á–∞ ‚Äî –∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å –∑–∞–ø–∏—Å–∏ –∏–∑ –¥–Ω–µ–≤–Ω–∏–∫–∞ –°–ú–≠–† –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –∏ –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–ª—è—Ç—å –æ–±–æ—Å–Ω–æ–≤–∞–Ω–Ω—ã–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏
                –ø–æ —É–ø—Ä–∞–≤–ª–µ–Ω–∏—é —ç–º–æ—Ü–∏—è–º–∏ –∏ —Ä–µ–∞–∫—Ü–∏—è–º–∏, –∞ —Ç–∞–∫–∂–µ –¥—Ä—É–≥–∏–µ –ø–æ–ª–µ–∑–Ω—ã–µ –ø—Å–∏—Ö–æ–ª–æ–≥–∏—á–µ—Å–∫–∏–µ —Å–æ–≤–µ—Ç—ã.
                –ü—Ä–∏ —ç—Ç–æ–º —Ç—ã –∏—Å–ø–æ–ª—å–∑—É–µ—à—å —Ç–æ–ª—å–∫–æ –¥–æ—Å—Ç–æ–≤–µ—Ä–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ –∏ –∏–∑–±–µ–≥–∞–µ—à—å –ª—é–±—ã—Ö –ø—Ä–µ–¥–ø–æ–ª–æ–∂–µ–Ω–∏–π –∏–ª–∏ –≤—ã–º—ã—à–ª–µ–Ω–Ω—ã—Ö —Ñ–∞–∫—Ç–æ–≤.
        
                –ü–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å –ø—Ä–µ–¥–æ—Å—Ç–∞–≤–∏–ª —Å–ª–µ–¥—É—é—â—É—é –∑–∞–ø–∏—Å—å:
        
                –°–∏—Ç—É–∞—Ü–∏—è: {situation}
                –ú—ã—Å–ª—å: {thought}
                –≠–º–æ—Ü–∏—è: {emotion}
                –†–µ–∞–∫—Ü–∏—è: {reaction}
        
                –ù–∞ –æ—Å–Ω–æ–≤–µ —ç—Ç–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏:
        
                1. –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä—É–π —Å–≤—è–∑—å –º–µ–∂–¥—É —Å–∏—Ç—É–∞—Ü–∏–µ–π, –º—ã—Å–ª—å—é, —ç–º–æ—Ü–∏–µ–π –∏ —Ä–µ–∞–∫—Ü–∏–µ–π, –≤—ã—è–≤–∏–≤ –≤–æ–∑–º–æ–∂–Ω—ã–µ –∫–æ–≥–Ω–∏—Ç–∏–≤–Ω—ã–µ –∏—Å–∫–∞–∂–µ–Ω–∏—è –∏–ª–∏ –ø–∞—Ç—Ç–µ—Ä–Ω—ã –ø–æ–≤–µ–¥–µ–Ω–∏—è.
                2. –ü—Ä–µ–¥–ª–æ–∂–∏ –∫–æ–Ω–∫—Ä–µ—Ç–Ω—ã–µ —Å—Ç—Ä–∞—Ç–µ–≥–∏–∏ –∏–ª–∏ —Ç–µ—Ö–Ω–∏–∫–∏ –¥–ª—è —É–ø—Ä–∞–≤–ª–µ–Ω–∏—è –¥–∞–Ω–Ω—ã–º–∏ —ç–º–æ—Ü–∏—è–º–∏ –∏ —Ä–µ–∞–∫—Ü–∏—è–º–∏, –æ–ø–∏—Ä–∞—è—Å—å –Ω–∞ –¥–æ–∫–∞–∑–∞–Ω–Ω—ã–µ –ø—Å–∏—Ö–æ–ª–æ–≥–∏—á–µ—Å–∫–∏–µ –º–µ—Ç–æ–¥—ã.
                3. –ò—Å–ø–æ–ª—å–∑—É–π –ø—Ä–∏–º–µ—Ä—ã –∏–∑ –∂–∏–∑–Ω–∏ –∏–ª–∏ –º–µ—Ç–∞—Ñ–æ—Ä—ã, —á—Ç–æ–±—ã –∏–ª–ª—é—Å—Ç—Ä–∏—Ä–æ–≤–∞—Ç—å –ø—Ä–µ–¥–ª–æ–∂–µ–Ω–Ω—ã–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –∏ —Å–¥–µ–ª–∞—Ç—å –∏—Ö –±–æ–ª–µ–µ –ø–æ–Ω—è—Ç–Ω—ã–º–∏ –∏ –ø—Ä–∏–º–µ–Ω–∏–º—ã–º–∏.
                4. –û–±—ä—è—Å–Ω–∏, –ø–æ—á–µ–º—É –∏–º–µ–Ω–Ω–æ —ç—Ç–∏ –ø–æ–¥—Ö–æ–¥—ã —ç—Ñ—Ñ–µ–∫—Ç–∏–≤–Ω—ã –≤ –¥–∞–Ω–Ω–æ–π —Å–∏—Ç—É–∞—Ü–∏–∏, —Å—Å—ã–ª–∞—è—Å—å –Ω–∞ –ø—Å–∏—Ö–æ–ª–æ–≥–∏—á–µ—Å–∫–∏–µ —Ç–µ–æ—Ä–∏–∏ –∏–ª–∏ –∏—Å—Å–ª–µ–¥–æ–≤–∞–Ω–∏—è.
                5. –î–æ–±–∞–≤—å –ø—Ä–∞–∫—Ç–∏—á–µ—Å–∫–∏–π —Å–æ–≤–µ—Ç –∏–ª–∏ —É–ø—Ä–∞–∂–Ω–µ–Ω–∏–µ, –∫–æ—Ç–æ—Ä–æ–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å —Å–º–æ–∂–µ—Ç –ª–µ–≥–∫–æ –≤–Ω–µ–¥—Ä–∏—Ç—å –≤ —Å–≤–æ—é –ø–æ–≤—Å–µ–¥–Ω–µ–≤–Ω—É—é –∂–∏–∑–Ω—å
                –¥–ª—è —É–ª—É—á—à–µ–Ω–∏—è —ç–º–æ—Ü–∏–æ–Ω–∞–ª—å–Ω–æ–≥–æ —Å–æ—Å—Ç–æ—è–Ω–∏—è –∏ —Ä–µ–∞–∫—Ü–∏–∏.
                
                –í–∞–∂–Ω–æ: –ù–µ –∏—Å–ø–æ–ª—å–∑—É–π —à–∞–±–ª–æ–Ω–Ω—ã–µ –∏–ª–∏ –æ–±—â–∏–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏.
                –î–∞–π —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –∏ —Å–æ–≤–µ—Ç—ã, –∞–¥–∞–ø—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ –∫ —ç—Ç–æ–π –∑–∞–ø–∏—Å–∏.
                –ù–µ –±–æ–ª–µ–µ 2000 —Å–∏–º–≤–æ–ª–æ–≤."""
        )
    ]

    try:
        # –ü–µ—Ä–µ–¥–∞–µ–º –≤ –≤—ã–∑–æ–≤ –∫–æ–Ω—Ñ–∏–≥—É—Ä–∞—Ü–∏—é —Å –Ω—É–∂–Ω—ã–º thread_id,
        # —á—Ç–æ –æ–±–µ—Å–ø–µ—á–∏–≤–∞–µ—Ç –ø–æ–¥–¥–µ—Ä–∂–∫—É –æ—Ç–¥–µ–ª—å–Ω—ã—Ö —Ä–∞–∑–≥–æ–≤–æ—Ä–æ–≤
        output = await app.ainvoke(
            {"messages": input_messages},
            config={"configurable": {"thread_id": user_id}}
        )
        analysis = output["messages"][-1].content

        # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –≤ –ë–î (–ø—Ä–∏–º–µ—Ä)
        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute(
                "UPDATE diary SET recommendation = ? WHERE id = ?",
                (analysis, entry_id)
            )
            await db.commit()

        await message.answer(analysis, reply_markup=dialog_buttons)

    except Exception as e:
        await message.answer(f"–û—à–∏–±–∫–∞ –∞–Ω–∞–ª–∏–∑–∞: {e}")

# –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –Ω–∞—Å—Ç—Ä–æ–µ–∫ (InlineKeyboard) –¥–∏–Ω–∞–º–∏—á–µ—Å–∫–∏
async def generate_settings_menu(user_id: int) -> InlineKeyboardMarkup:
    async with aiosqlite.connect(DB_PATH) as db:
        async with db.execute(
            "SELECT enabled FROM reminders WHERE user_id = ?", (user_id,)
        ) as cursor:
            result = await cursor.fetchone()
            reminders_enabled = result[0] if result else 0

    buttons = [
        [
            InlineKeyboardButton(
                text="\u2705 –ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è: –í–∫–ª—é—á–∏—Ç—å" if not reminders_enabled else "\u274C –ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è: –í—ã–∫–ª—é—á–∏—Ç—å",
                callback_data="toggle_reminder_on" if not reminders_enabled else "toggle_reminder_off"
            )
        ]
    ]

    if reminders_enabled:
        buttons.append([
            InlineKeyboardButton(
                text="–£—Å—Ç–∞–Ω–æ–≤–∏—Ç—å –≤—Ä–µ–º—è –Ω–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è",
                callback_data="set_reminder_time"
            )
        ])

    return InlineKeyboardMarkup(inline_keyboard=buttons)

@dp.callback_query(lambda c: c.data == "continue_dialog")
async def continue_dialog(callback_query: CallbackQuery, state: FSMContext):
    await callback_query.message.answer("–ü—Ä–æ–¥–æ–ª–∂–∞–π—Ç–µ –¥–∏–∞–ª–æ–≥. –ù–∞–ø–∏—à–∏—Ç–µ –≤–∞—à –≤–æ–ø—Ä–æ—Å:", reply_markup=ReplyKeyboardRemove())
    user_id = str(callback_query.from_user.id)
    await state.set_state(DialogForm.in_dialog)
    await callback_query.answer()

@dp.message(DialogForm.in_dialog)
async def dialog_interaction(message: Message):
    input_messages = [HumanMessage(content=message.text)]
    user_id = str(message.from_user.id)
    
    try:
        output = await app.ainvoke(
            {"messages": input_messages},
            config={"configurable": {"thread_id": user_id}}
        )
        response = output["messages"][-1].content
        await message.answer(response, reply_markup=dialog_buttons)
    except Exception as e:
        await message.answer(f"–û—à–∏–±–∫–∞ –≤–æ –≤—Ä–µ–º—è –¥–∏–∞–ª–æ–≥–∞: {e}")

@dp.callback_query(lambda c: c.data == "end_dialog")
async def end_dialog(callback_query: CallbackQuery, state: FSMContext):
    user_id = str(callback_query.from_user.id)

    # –£–±–∏—Ä–∞–µ–º inline-–∫–Ω–æ–ø–∫–∏
    await callback_query.message.edit_reply_markup(reply_markup=None)

    # –û—á–∏—â–∞–µ–º FSM —Å–æ—Å—Ç–æ—è–Ω–∏–µ
    await state.clear()

    # –£–≤–µ–¥–æ–º–ª—è–µ–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –æ –∑–∞–≤–µ—Ä—à–µ–Ω–∏–∏
    await callback_query.answer("–î–∏–∞–ª–æ–≥ –∑–∞–≤–µ—Ä—à—ë–Ω.")
    await callback_query.message.answer(
        "–î–∏–∞–ª–æ–≥ –∑–∞–≤–µ—Ä—à—ë–Ω. –ï—Å–ª–∏ –ø–æ–Ω–∞–¥–æ–±–∏—Ç—Å—è –ø–æ–º–æ—â—å —Å–Ω–æ–≤–∞, –≤—ã–±–µ—Ä–∏—Ç–µ –Ω—É–∂–Ω—ã–π –ø—É–Ω–∫—Ç –º–µ–Ω—é.",
        reply_markup=main_menu
    )

@dp.message(Command(commands=["help"]))
async def cmd_help(message: Message):
    help_text = (
        "–ü—Ä–∏–≤–µ—Ç! üëã –Ø ‚Äî —Ç–≤–æ–π –ª–∏—á–Ω—ã–π –¥–Ω–µ–≤–Ω–∏–∫ –∏ –ø–æ–º–æ—â–Ω–∏–∫ –ø–æ –ø—Å–∏—Ö–∏—á–µ—Å–∫–æ–º—É –∑–¥–æ—Ä–æ–≤—å—é.\n"
        "–Ø –ø–æ–º–æ–≥—É —Ñ–∏–∫—Å–∏—Ä–æ–≤–∞—Ç—å —Å–æ–±—ã—Ç–∏—è, –º—ã—Å–ª–∏ –∏ —ç–º–æ—Ü–∏–∏, –∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å –∏—Ö –∏ –ø–æ–ª—É—á–∞—Ç—å –ø–æ–ª–µ–∑–Ω—ã–µ —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏ –ø–æ –ö–ü–¢.\n\n"
        "‚ú® –û—Å–Ω–æ–≤–Ω—ã–µ –∫–æ–º–∞–Ω–¥—ã:\n"
        "/start ‚Äî –Ω–∞—á–∞—Ç—å —Ä–∞–±–æ—Ç—É –∏ –∑–∞—Ä–µ–≥–∏—Å—Ç—Ä–∏—Ä–æ–≤–∞—Ç—å—Å—è.\n"
        "/new_entry ‚Äî –¥–æ–±–∞–≤–∏—Ç—å –Ω–æ–≤—É—é –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫.\n"
        "/get_recommendation ‚Äî –ø–æ–ª—É—á–∏—Ç—å —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏–∏.\n"
        "/view_diary ‚Äî –ø–æ—Å–º–æ—Ç—Ä–µ—Ç—å –¥–Ω–µ–≤–Ω–∏–∫.\n"
        "/export_diary ‚Äî —Å–∫–∞—á–∞—Ç—å –¥–Ω–µ–≤–Ω–∏–∫.\n"
        "/feedback ‚Äî –æ—Å—Ç–∞–≤–∏—Ç—å –æ—Ç–∑—ã–≤.\n"
        "/settings ‚Äî –æ—Ç–∫—Ä—ã—Ç—å –Ω–∞—Å—Ç—Ä–æ–π–∫–∏.\n"
        "/help ‚Äî –æ–ø–∏—Å–∞–Ω–∏–µ –∫–æ–º–∞–Ω–¥.\n\n"

        "üóÇ –ì–ª–∞–≤–Ω–æ–µ –º–µ–Ω—é:\n"
        "üîπ –î–æ–±–∞–≤–∏—Ç—å –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫ ‚Äì –∑–∞–ø–∏—Å—ã–≤–∞–π –º—ã—Å–ª–∏ –∏ –ø–µ—Ä–µ–∂–∏–≤–∞–Ω–∏—è.\n"
        "üîπ –ü–æ–ª—É—á–∏—Ç—å —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—é ‚Äì –∞–Ω–∞–ª–∏–∑–∏—Ä—É–π —Å–≤–æ—é –ø–æ—Å–ª–µ–¥–Ω—é—é –∑–∞–ø–∏—Å—å —Å –ø–æ–º–æ—â—å—é –∏—Å–∫—É—Å—Å—Ç–≤–µ–Ω–Ω–æ–≥–æ –∏–Ω—Ç–µ–ª–ª–µ–∫—Ç–∞.\n"
        "üîπ –ü–æ—Å–º–æ—Ç—Ä–µ—Ç—å –¥–Ω–µ–≤–Ω–∏–∫ ‚Äì –ø–µ—Ä–µ—á–∏—Ç—ã–≤–∞–π –∑–∞–ø–∏—Å–∏ –∏ —Å–ª–µ–¥–∏ –∑–∞ –ø—Ä–æ–≥—Ä–µ—Å—Å–æ–º.\n"
        "üîπ –≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –¥–Ω–µ–≤–Ω–∏–∫ ‚Äì –∑–∞–≥—Ä—É–∂–∞–π –∑–∞–ø–∏—Å–∏ –≤ —É–¥–æ–±–Ω–æ–º —Ñ–æ—Ä–º–∞—Ç–µ.\n"
        "üîπ –û—Å—Ç–∞–≤–∏—Ç—å –æ—Ç–∑—ã–≤ ‚Äì –ø–æ–º–æ–≥–∏ —Å–¥–µ–ª–∞—Ç—å –±–æ—Ç–∞ –ª—É—á—à–µ.\n"
        "üîπ –ù–∞—Å—Ç—Ä–æ–π–∫–∏ ‚Äì –Ω–∞—Å—Ç—Ä–æ–π —É–≤–µ–¥–æ–º–ª–µ–Ω–∏—è, —á—Ç–æ–±—ã –Ω–µ –∑–∞–±—ã—Ç—å –∑–∞–ø–æ–ª–Ω–∏—Ç—å –¥–Ω–µ–≤–Ω–∏–∫ (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é 18:00, –µ—Å–ª–∏ –Ω–µ —É–∫–∞–∑—ã–≤–∞–ª–∏ —Ä–∞–Ω–µ–µ)"
    )
    await message.answer(help_text)

@dp.message(lambda m: m.text == "–≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –¥–Ω–µ–≤–Ω–∏–∫")
@dp.message(Command(commands=["export_diary"]))
async def handle_export_diary(message: Message):
    user_id = message.from_user.id
    
    # –ü–æ–ª—É—á–µ–Ω–∏–µ –≤—Å–µ—Ö –∑–∞–ø–∏—Å–µ–π –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    async with aiosqlite.connect(DB_PATH) as db:
        async with db.execute(
            """
            SELECT situation, thought, emotion, reaction, recommendation, created_at
            FROM diary
            WHERE user_id = ?
            ORDER BY created_at ASC
            """,
            (user_id,)
        ) as cursor:
            entries = await cursor.fetchall()

    if not entries:
        await message.answer("–í–∞—à –¥–Ω–µ–≤–Ω–∏–∫ –ø—É—Å—Ç. –î–æ–±–∞–≤—å—Ç–µ –∑–∞–ø–∏—Å–∏, —á—Ç–æ–±—ã —ç–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –∏—Ö.")
        return

    # –ì–µ–Ω–µ—Ä–∞—Ü–∏—è DOCX-—Ñ–∞–π–ª–∞
    document = Document()
    document.add_heading("–î–Ω–µ–≤–Ω–∏–∫ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è", level=1)

    for idx, (situation, thought, emotion, reaction, recommendation, created_at) in enumerate(entries, start=1):
        document.add_heading(f"–ó–∞–ø–∏—Å—å #{idx}", level=2)
        document.add_paragraph(f"–î–∞—Ç–∞: {created_at}")
        document.add_paragraph(f"–°–∏—Ç—É–∞—Ü–∏—è: {situation}")
        document.add_paragraph(f"–ú—ã—Å–ª—å: {thought}")
        document.add_paragraph(f"–≠–º–æ—Ü–∏—è: {emotion}")
        document.add_paragraph(f"–†–µ–∞–∫—Ü–∏—è: {reaction}")
        document.add_paragraph(f"–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—è: {recommendation or '–ù–µ –ø–æ–ª—É—á–µ–Ω–∞'}")
        document.add_paragraph("-" * 118)

    file_path = f"tmp/diary_{user_id}.docx"
    document.save(file_path)

    input_file = FSInputFile(file_path)
    await message.answer_document(
        document=input_file,
        caption="–í–∞—à –¥–Ω–µ–≤–Ω–∏–∫ –≤ —Ñ–æ—Ä–º–∞—Ç–µ DOCX"
    )
    os.remove(file_path)

@dp.message(lambda m: m.text == "–ü–æ—Å–º–æ—Ç—Ä–µ—Ç—å –¥–Ω–µ–≤–Ω–∏–∫")
@dp.message(Command(commands=["view_diary"]))
async def handle_view_diary(message: Message):
    user_id = message.from_user.id

    # –ü–æ–ª—É—á–µ–Ω–∏–µ –≤—Å–µ—Ö –∑–∞–ø–∏—Å–µ–π –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è
    async with aiosqlite.connect(DB_PATH) as db:
        async with db.execute(
            """
            SELECT id, situation, thought, emotion, reaction, recommendation, created_at
            FROM diary
            WHERE user_id = ?
            ORDER BY created_at ASC
            """,
            (user_id,)
        ) as cursor:
            entries = await cursor.fetchall()

    if not entries:
        await message.answer("–í–∞—à –¥–Ω–µ–≤–Ω–∏–∫ –ø—É—Å—Ç. –î–æ–±–∞–≤—å—Ç–µ –∑–∞–ø–∏—Å–∏, —á—Ç–æ–±—ã –∏—Ö —É–≤–∏–¥–µ—Ç—å.")
        return

    for entry in entries:
        entry_id, situation, thought, emotion, reaction, recommendation, created_at = entry

        # –§–æ—Ä–º–∏—Ä—É–µ–º —Ç–µ–∫—Å—Ç –∑–∞–ø–∏—Å–∏
        diary_text = (
            f"<b>–î–∞—Ç–∞:</b> {created_at}\n"
            f"<b>–°–∏—Ç—É–∞—Ü–∏—è:</b> {situation}\n"
            f"<b>–ú—ã—Å–ª—å:</b> {thought}\n"
            f"<b>–≠–º–æ—Ü–∏—è:</b> {emotion}\n"
            f"<b>–†–µ–∞–∫—Ü–∏—è:</b> {reaction}\n"
            f"<b>–†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—è:</b> {recommendation or '–ù–µ –ø–æ–ª—É—á–µ–Ω–∞'}"
        )

        # –°–æ–∑–¥–∞–µ–º inline-–∫–Ω–æ–ø–∫—É –¥–ª—è —É–¥–∞–ª–µ–Ω–∏—è –∑–∞–ø–∏—Å–∏
        delete_button = InlineKeyboardMarkup(inline_keyboard=[
            [
                InlineKeyboardButton(
                    text="–£–¥–∞–ª–∏—Ç—å –∑–∞–ø–∏—Å—å",
                    callback_data=f"delete_diary_{entry_id}"
                )
            ]
        ])

        await message.answer(diary_text, reply_markup=delete_button, parse_mode=ParseMode.HTML)

@dp.callback_query(lambda c: c.data.startswith("delete_diary_"))
async def handle_delete_diary(callback_query: CallbackQuery):
    entry_id = int(callback_query.data.split("_")[2])

    # –£–¥–∞–ª—è–µ–º –∑–∞–ø–∏—Å—å –∏–∑ –±–∞–∑—ã –¥–∞–Ω–Ω—ã—Ö
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            "DELETE FROM diary WHERE id = ?",
            (entry_id,)
        )
        await db.commit()

    # –£–≤–µ–¥–æ–º–ª—è–µ–º –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—è –æ–± —É—Å–ø–µ—à–Ω–æ–º —É–¥–∞–ª–µ–Ω–∏–∏
    await callback_query.answer("–ó–∞–ø–∏—Å—å —É—Å–ø–µ—à–Ω–æ —É–¥–∞–ª–µ–Ω–∞!")

    # –û–±–Ω–æ–≤–ª—è–µ–º —Å–æ–æ–±—â–µ–Ω–∏–µ, —É–±–∏—Ä–∞—è inline-–∫–Ω–æ–ø–∫–∏
    await callback_query.message.edit_reply_markup(reply_markup=None)
    await callback_query.message.answer("–ó–∞–ø–∏—Å—å –±—ã–ª–∞ —É–¥–∞–ª–µ–Ω–∞.")

@dp.message(lambda m: m.text == "–ù–∞—Å—Ç—Ä–æ–π–∫–∏")
@dp.message(Command(commands=["settings"]))
async def handle_menu_settings(message: Message):
    user_id = message.from_user.id
    settings_menu = await generate_settings_menu(user_id)
    await message.answer("–ù–∞—Å—Ç—Ä–æ–π–∫–∏ –Ω–∞–ø–æ–º–∏–Ω–∞–Ω–∏–π:", reply_markup=settings_menu)

@dp.callback_query(lambda c: c.data in ["toggle_reminder_on", "toggle_reminder_off"])
async def toggle_reminders(callback_query: CallbackQuery):
    user_id = callback_query.from_user.id
    enable_reminders = callback_query.data == "toggle_reminder_on"

    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            """
            INSERT OR REPLACE INTO reminders (user_id, enabled)
            VALUES (?, ?)
            """,
            (user_id, 1 if enable_reminders else 0)
        )
        await db.commit()

    await callback_query.answer(
        "–ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è –≤–∫–ª—é—á–µ–Ω—ã!" if enable_reminders else "–ù–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è –≤—ã–∫–ª—é—á–µ–Ω—ã!"
    )

    settings_menu = await generate_settings_menu(user_id)
    await callback_query.message.edit_reply_markup(reply_markup=settings_menu)

@dp.callback_query(lambda c: c.data == "set_reminder_time")
async def set_reminder_time(callback_query: CallbackQuery, state: FSMContext):
    await state.set_state(ReminderForm.time)
    await callback_query.message.answer(
        "–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, —É–∫–∞–∂–∏—Ç–µ –≤—Ä–µ–º—è –¥–ª—è –Ω–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è –≤ —Ñ–æ—Ä–º–∞—Ç–µ –ß–ß:–ú–ú (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é 18:00)",
        reply_markup=ReplyKeyboardRemove()
    )
    await callback_query.answer()

@dp.message(ReminderForm.time)
async def process_reminder_time(message: Message, state: FSMContext):
    try:
        reminder_time = message.text if message.text else "18:00"
        if len(reminder_time) != 5 or reminder_time[2] != ":" or not reminder_time.replace(":", "").isdigit():
            raise ValueError("–ù–µ–≤–µ—Ä–Ω—ã–π —Ñ–æ—Ä–º–∞—Ç –≤—Ä–µ–º–µ–Ω–∏")

        hours, minutes = map(int, reminder_time.split(":"))
        if not (0 <= hours < 24 and 0 <= minutes < 60):
            raise ValueError("–ß–∞—Å—ã –∏–ª–∏ –º–∏–Ω—É—Ç—ã –≤—ã—Ö–æ–¥—è—Ç –∑–∞ –¥–æ–ø—É—Å—Ç–∏–º—ã–π –¥–∏–∞–ø–∞–∑–æ–Ω")

        user_id = message.from_user.id
        async with aiosqlite.connect(DB_PATH) as db:
            await db.execute(
                """
                INSERT OR REPLACE INTO reminders (user_id, enabled, time)
                VALUES (?, 1, ?)
                """,
                (user_id, reminder_time)
            )
            await db.commit()

        await state.clear()
        await message.answer(
            f"–í—Ä–µ–º—è –Ω–∞–ø–æ–º–∏–Ω–∞–Ω–∏—è —É—Å–ø–µ—à–Ω–æ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω–æ –Ω–∞ {reminder_time}!",
            reply_markup=main_menu
        )

    except ValueError as e:
        await message.answer(f"–û—à–∏–±–∫–∞: {e}. –ü–æ–ø—Ä–æ–±—É–π—Ç–µ –µ—â—ë —Ä–∞–∑ –≤ —Ñ–æ—Ä–º–∞—Ç–µ –ß–ß:–ú–ú.")

# –ê–Ω–∫–µ—Ç–∞ –æ–±—Ä–∞—Ç–Ω–æ–π —Å–≤—è–∑–∏
class FeedbackForm(StatesGroup):
    feedback = State()

@dp.message(lambda m: m.text == "–û—Å—Ç–∞–≤–∏—Ç—å –æ—Ç–∑—ã–≤")
@dp.message(Command(commands=["feedback"]))
async def handle_menu_feedback(message: Message, state: FSMContext):
    await state.set_state(FeedbackForm.feedback)
    await message.answer("–ü–æ–∂–∞–ª—É–π—Å—Ç–∞, –æ—Å—Ç–∞–≤—å—Ç–µ –≤–∞—à –æ—Ç–∑—ã–≤:", reply_markup=ReplyKeyboardRemove())

@dp.message(FeedbackForm.feedback)
async def process_feedback(message: Message, state: FSMContext):
    feedback = message.text
    user_id = message.from_user.id

    # –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –æ—Ç–∑—ã–≤–∞ –≤ –ë–î
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            """
            INSERT INTO feedback (user_id, feedback)
            VALUES (?, ?)
            """,
            (user_id, feedback)
        )
        await db.commit()

    await state.clear()
    await message.answer("–°–ø–∞—Å–∏–±–æ –∑–∞ –≤–∞—à –æ—Ç–∑—ã–≤!", reply_markup=main_menu)

@dp.message(lambda message: message.text.lower() in ["–≥–æ–¥–∂–æ —Å–∞—Ç–æ—Ä—É", "gojo satoru", "gojo", "satoru gojo",
                                                     "—Å–∞—Ç–æ—Ä—É –≥–æ–¥–∂–æ", "–≥–æ–¥–∂–æ", "—Å–∞—Ç–æ—Ä—É", "satoru",
                                                     "–≥–æ–¥–∂–æ —Å–∞—Ç–æ—Ä–∏", "—Å–∞—Ç–æ—Ä–∏ –≥–æ–¥–∂–æ", "—Å–∞—Ç–æ—Ä–∏"])
async def send_gojo_image(message: Message):
    try:
        # –ü—É—Ç—å –∫ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—é
        image_path = "pic/gojo.jpg"
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞–ª–∏—á–∏—è —Ñ–∞–π–ª–∞
        if not os.path.exists(image_path):
            await message.answer("–≠—Ö, –ø–∞—Å—Ö–∞–ª–∫–∞ –ø–æ—Ç–µ—Ä—è–ª–∞—Å—å")
            return

        # –û—Ç–ø—Ä–∞–≤–∫–∞ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—é
        photo = FSInputFile(image_path)
        await message.answer_photo(photo, caption="–í–æ—Ç –æ–Ω –º–∏—Å—Ç–µ—Ä –ø–∞—Å—Ö–∞–ª–∫–∏–Ω!")

    except Exception as e:
        await message.answer(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ—Ç–ø—Ä–∞–≤–∫–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è: {e}")



@dp.message()
async def unknown_message(message: Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state in [DiaryForm.situation, DiaryForm.thought, DiaryForm.emotion,
                         DiaryForm.reaction, ReminderForm.time, DialogForm.in_dialog,
                         FeedbackForm.feedback, ReminderForm.time]:
        return

    await message.answer(
        "–û–π, –∫–∞–∂–µ—Ç—Å—è –≤—ã –ø–æ–ø–∞–ª–∏ –≤ –Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–æ–µ –º–µ—Å—Ç–æ, –ø–æ–ø—Ä–æ–±—É–π—Ç–µ –¥—Ä—É–≥—É—é –∫–æ–º–∞–Ω–¥—É –∏–ª–∏ –∫–Ω–æ–ø–∫—É :)",
        reply_markup=main_menu
    )
# ------------------------------------------------------------------------------
# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è –±–∞–∑—ã –¥–∞–Ω–Ω—ã—Ö
# ------------------------------------------------------------------------------
async def init_db():
    async with aiosqlite.connect(DB_PATH) as db:
        await db.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY,
                name TEXT,
                age INTEGER,
                email TEXT,
                created_at DATETIME DEFAULT (DATETIME('now', '+3 hours'))
            )
            """
        )
        await db.execute(
            """
            CREATE TABLE IF NOT EXISTS diary (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                situation TEXT,
                thought TEXT,
                emotion TEXT,
                reaction TEXT,
                recommendation TEXT DEFAULT NULL,
                created_at DATETIME DEFAULT (DATETIME('now', '+3 hours')),
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
            """
        )
        await db.execute(
            """
            CREATE TABLE IF NOT EXISTS reminders (
                user_id INTEGER PRIMARY KEY,
                enabled INTEGER DEFAULT 0,
                time TEXT DEFAULT '18:00',
                last_sent_date DATE,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
            """
        )
        await db.execute(
            """
            CREATE TABLE IF NOT EXISTS feedback (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                feedback TEXT,
                created_at DATETIME DEFAULT (DATETIME('now', '+3 hours')),
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
            """
        )
        await db.commit()

# ------------------------------------------------------------------------------
# Main Entry Point
# ------------------------------------------------------------------------------
async def main():
    # –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö, –∫–æ–º–∞–Ω–¥—ã –±–æ—Ç–∞, –ø–ª–∞–Ω–∏—Ä–æ–≤—â–∏–∫ –∏ —Ç.–¥.
    await init_db()
    await bot.set_my_commands([
        BotCommand(command="start", description="–ù–∞—á–∞—Ç—å —Ä–∞–±–æ—Ç—É"),
        BotCommand(command="new_entry", description="–î–æ–±–∞–≤–∏—Ç—å –∑–∞–ø–∏—Å—å –≤ –¥–Ω–µ–≤–Ω–∏–∫"),
        BotCommand(command="get_recommendation", description="–ü–æ–ª—É—á–∏—Ç—å —Ä–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—é"),
        BotCommand(command="view_diary", description="–ü–æ—Å–º–æ—Ç—Ä–µ—Ç—å –¥–Ω–µ–≤–Ω–∏–∫"),
        BotCommand(command="export_diary", description="–≠–∫—Å–ø–æ—Ä—Ç–∏—Ä–æ–≤–∞—Ç—å –¥–Ω–µ–≤–Ω–∏–∫"),
        BotCommand(command="feedback", description="–û—Å—Ç–∞–≤–∏—Ç—å –æ—Ç–∑—ã–≤"),
        BotCommand(command="settings", description="–û—Ç–∫—Ä—ã—Ç—å –Ω–∞—Å—Ç—Ä–æ–π–∫–∏"),
        BotCommand(command="help", description="–ü–æ–ª–Ω–æ–µ –æ–ø–∏—Å–∞–Ω–∏–µ –∫–æ–º–∞–Ω–¥ –∏ –∫–Ω–æ–ø–æ–∫ –±–æ—Ç–∞")
    ])
    if not scheduler.get_jobs():
        scheduler.remove_all_jobs()
        scheduler.add_job(send_reminders, "cron", second=0)
    if not scheduler.running:
        scheduler.start()
    await dp.start_polling(bot, timeout=60)

if __name__ == "__main__":
    try:
        nest_asyncio.apply()
        asyncio.run(main())
    except (KeyboardInterrupt, SystemExit):
        logger.error("–ë–æ—Ç –æ—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω!")